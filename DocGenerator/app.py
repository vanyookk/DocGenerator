from flask import Flask, request, send_file
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
import datetime
import os
from pathlib import Path
import logging

app = Flask(__name__)

# Настройка логирования
logging.basicConfig(filename='document_errors.log', level=logging.ERROR)

# HTML-шаблон с улучшенным дизайном
HTML_TEMPLATE = """
<!DOCTYPE html>
<html lang="ru">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Генератор документов</title>
    <link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600&display=swap" rel="stylesheet">
    <style>
        :root {
            --bg-1: #f8f9ff;
            --bg-2: #eef0ff;
            --card-bg: rgba(255, 255, 255, 0.95);
            --primary: #6366f1;
            --primary-hover: #4f46e5;
            --text: #1e293b;
            --border: #e2e8f0;
            --shadow: 0 12px 32px rgba(0,0,0,0.08);
            --radius: 16px;
        }

        * {
            box-sizing: border-box;
            margin: 0;
            padding: 0;
            font-family: 'Inter', system-ui, -apple-system, sans-serif;
        }

        body {
            min-height: 100vh;
            display: flex;
            align-items: center;
            justify-content: center;
            background: linear-gradient(45deg, var(--bg-1), var(--bg-2));
            color: var(--text);
            position: relative;
            overflow: hidden;
        }

        body::before {
            content: '';
            position: absolute;
            width: 150%;
            height: 150%;
            background: radial-gradient(circle at 50% 50%, 
                rgba(199, 210, 254, 0.15) 0%, 
                rgba(199, 210, 254, 0) 60%);
            animation: gradient-pulse 20s infinite alternate;
        }

        .container {
            background: var(--card-bg);
            backdrop-filter: blur(12px);
            border-radius: var(--radius);
            box-shadow: var(--shadow);
            padding: 2.5rem;
            width: 100%;
            max-width: 480px;
            margin: 1rem;
            border: 1px solid rgba(255,255,255,0.3);
            transform: translateY(0);
            transition: transform 0.3s ease, box-shadow 0.3s ease;
        }

        .container:hover {
            transform: translateY(-4px);
            box-shadow: 0 24px 48px rgba(0,0,0,0.1);
        }

        h1 {
            font-size: 1.8rem;
            font-weight: 600;
            margin-bottom: 2rem;
            text-align: center;
            background: linear-gradient(45deg, var(--primary), #8b5cf6);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
        }

        .form-group {
            margin-bottom: 1.5rem;
        }

        label {
            display: block;
            margin-bottom: 0.5rem;
            font-size: 0.9rem;
            font-weight: 500;
            color: var(--text);
            opacity: 0.9;
        }

        input, select {
            width: 100%;
            padding: 0.9rem;
            border: 2px solid var(--border);
            border-radius: 10px;
            font-size: 1rem;
            transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
            background: rgba(255,255,255,0.8);
        }

        input:focus, select:focus {
            outline: none;
            border-color: var(--primary);
            box-shadow: 0 0 0 4px rgba(99, 102, 241, 0.15);
            transform: scale(1.02);
        }

        input:hover, select:hover {
            border-color: #c7d2fe;
        }

        button {
            width: 100%;
            padding: 1rem;
            background: var(--primary);
            color: white;
            border: none;
            border-radius: 10px;
            font-size: 1rem;
            font-weight: 600;
            cursor: pointer;
            transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
            position: relative;
            overflow: hidden;
        }

        button::before {
            content: '';
            position: absolute;
            top: 0;
            left: -100%;
            width: 100%;
            height: 100%;
            background: linear-gradient(
                120deg,
                transparent,
                rgba(255,255,255,0.3),
                transparent
            );
            transition: 0.5s;
        }

        button:hover {
            background: var(--primary-hover);
            transform: translateY(-2px);
            box-shadow: 0 8px 16px rgba(99, 102, 241, 0.2);
        }

        button:hover::before {
            left: 100%;
        }

        .error {
            color: #ef4444;
            font-size: 0.9rem;
            text-align: center;
            margin-top: 1.5rem;
            opacity: 0;
            animation: fadeIn 0.3s ease forwards;
        }

        @keyframes gradient-pulse {
            0% { transform: rotate(0deg) scale(1); }
            100% { transform: rotate(5deg) scale(1.1); }
        }

        @keyframes fadeIn {
            from { opacity: 0; transform: translateY(10px); }
            to { opacity: 1; transform: translateY(0); }
        }
    </style>
</head>
<body>
    <div class="container">
        <h1>Генератор Документов</h1>
        <form method="POST">
            <div class="form-group">
                <label>Тип документа</label>
                <select name="doc_type">
                    <option value="dogovor">Договор</option>
                    <option value="schet">Счёт</option>
                    <option value="zayavlenie">Заявление</option>
                </select>
            </div>
            
            <div class="form-group">
                <label>Номер документа</label>
                <input type="text" name="number" required>
            </div>
            
            <div class="form-group">
                <label>Ваша организация</label>
                <input type="text" name="company_name" required>
            </div>
            
            <div class="form-group">
                <label>Клиент/Получатель</label>
                <input type="text" name="client_name" required>
            </div>
            
            <div class="form-group">
                <label>Город</label>
                <input type="text" name="city" value="Москва">
            </div>
            
            <div class="form-group">
                <label>Сумма (руб)</label>
                <input type="number" name="amount">
            </div>
            
            <button type="submit">
                Создать документ → 
            </button>
            
            {% if error %}
            <div class="error">{{ error }}</div>
            {% endif %}
        </form>
    </div>
</body>
</html>
"""
def create_document(doc_type, data):
    try:
        # Проверка обязательных полей
        required_fields = ['number', 'company_name', 'client_name']
        for field in required_fields:
            if field not in data:
                raise ValueError(f"Отсутствует обязательное поле: {field}")

        doc = Document()
        
        # Папки для документов
        folders = {
            "dogovor": "dogovor",
            "schet": "schet",
            "zayavlenie": "zayavlenie"
        }
        
        # Проверка типа документа
        if doc_type not in folders:
            raise ValueError(f"Неизвестный тип документа: {doc_type}")

        folder = folders[doc_type]
        Path(folder).mkdir(exist_ok=True)  # Создание папки если нужно

        # Стили документа
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)

        # Заголовок
        title = doc.add_paragraph(style='Heading 1')
        doc_type_russian = {
            "dogovor": "ДОГОВОР",
            "schet": "СЧЕТ",
            "zayavlenie": "ЗАЯВЛЕНИЕ"
        }.get(doc_type, doc_type.upper())
        
        title_run = title.add_run(f"{doc_type_russian} № {data['number']}")
        title_run.font.size = Pt(14)
        title_run.bold = True
        title.alignment = 1  # CENTER
        
        # Основное содержимое
        doc.add_paragraph(f"Дата: {data.get('date', datetime.datetime.now().strftime('%d.%m.%Y'))}")
        doc.add_paragraph(f"От: {data['company_name']}")
        doc.add_paragraph(f"Для: {data['client_name']}")
        doc.add_paragraph()  # Пустая строка

        # Блоки специфичные для каждого типа
        content_blocks = {
            "schet": [
                ("Сумма", f"{data.get('amount', 'Не указана')} руб."),
                ("Условия оплаты", data.get('payment_terms', 'По договорённости')),
                ("Реквизиты", data.get('details', 'Не указаны'))
            ],
            "dogovor": [
                ("Предмет договора", data.get('subject', 'Не указан')),
                ("Срок действия", data.get('validity_period', 'Не ограничен')),
                ("Особые условия", data.get('special_terms', 'Отсутствуют'))
            ],
            "zayavlenie": [
                ("Тип заявления", data.get('application_type', 'Не указан')),
                ("Содержание", data.get('content', 'Текст не предоставлен')),
                ("Основание", data.get('reason', 'Не указано'))
            ]
        }

        for label, value in content_blocks.get(doc_type, []):
            doc.add_paragraph(f"{label}: {value}")

        # Подписи
        doc.add_paragraph()  # Пустая строка
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        row = table.rows[0]
        row.cells[0].text = "Подпись:\n\n_________________"
        row.cells[1].text = "Дата:\n\n_________________"

        # Генерация имени файла
        filename = Path(folder) / f"{doc_type}_{data['number']}_{data.get('date', datetime.datetime.now().strftime('%d.%m.%Y'))}.docx"
        
        # Проверка доступности файла для записи
        try:
            doc.save(filename)
            return str(filename)
        except PermissionError:
            error_msg = f"Ошибка доступа к файлу: {filename}. Возможно, файл открыт в другой программе."
            logging.error(error_msg)
            raise PermissionError(error_msg)

    except Exception as e:
        error_msg = f"Ошибка при создании документа: {str(e)}"
        logging.error(error_msg)
        raise  # Пробрасываем исключение дальше для обработки в веб-приложении

@app.route("/", methods=["GET", "POST"])
def home():
    error = None
    if request.method == "POST":
        try:
            data = {
                "doc_type": request.form.get("doc_type", "dogovor"),
                "number": request.form.get("number", ""),
                "company_name": request.form.get("company_name", ""),
                "client_name": request.form.get("client_name", ""),
                "city": request.form.get("city", "Москва"),
                "amount": request.form.get("amount", ""),
                "date": datetime.datetime.now().strftime("%d.%m.%Y"),
                # Дополнительные поля для разных типов документов
                "payment_terms": request.form.get("payment_terms", "Безналичный расчет, 100% предоплата"),
                "subject": request.form.get("subject", "Оказание услуг"),
                "application_type": request.form.get("application_type", "Стандартное")
            }
            
            if not all([data["number"], data["company_name"], data["client_name"]]):
                error = "Пожалуйста, заполните все обязательные поля"
            else:
                filename = create_document(data["doc_type"], data)
                if filename and os.path.exists(filename):
                    return send_file(filename, as_attachment=True)
                error = "Ошибка при создании документа"
                
        except Exception as e:
            error = f"Произошла ошибка: {str(e)}"
            logging.error(f"Ошибка в обработке запроса: {str(e)}")
    
    return HTML_TEMPLATE.replace("{% if error %}", "").replace("{% endif %}", "").replace("{{ error }}", error if error else "")

if __name__ == "__main__":
    # Создаем папки для документов при запуске
    for folder in ["dogovor", "schet", "zayavlenie"]:
        os.makedirs(folder, exist_ok=True)
    app.run(debug=True, port=5000)